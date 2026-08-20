import json, re
from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).parent
OUT = ROOT / "Shopify_Theme_Editor_Management_Guide.docx"
TEMPLATES = [
    ("Homepage", "templates/index.json"),
    ("Product Page", "templates/product.json"),
    ("Collection Page", "templates/collection.json"),
]

def load_template(rel):
    s=(ROOT/rel).read_text(encoding="utf-8-sig")
    return json.loads(re.sub(r"/\*.*?\*/", "", s, flags=re.S))

def schema_for(path):
    if not path.exists(): return {}
    s=path.read_text(encoding="utf-8-sig")
    m=re.search(r"{%\s*schema\s*%}(.*?){%\s*endschema\s*%}", s, re.S)
    if not m: return {}
    try: return json.loads(m.group(1))
    except Exception: return {}

def title(v):
    if not v: return "Unnamed"
    if isinstance(v, str) and v.startswith("t:"): v=v.split(".")[-1]
    return str(v).replace("_", " ").replace("-", " ").strip().title()

def block_file(t):
    if t.startswith("shopify://apps/"): return None
    p=ROOT/"blocks"/(t+".liquid")
    return p if p.exists() else None

def section_file(t):
    p=ROOT/"sections"/(t+".liquid")
    return p if p.exists() else None

def fmt(v):
    if v is None: return "Not set"
    if v is True: return "Enabled"
    if v is False: return "Disabled"
    if isinstance(v,(dict,list)): return json.dumps(v, ensure_ascii=False)
    if v == "": return "Blank"
    return str(v).replace("{{ settings.color_palette.", "Global palette: ").replace(" }}", "")

PURPOSES={
"hero-slider-horizon":"Full-width rotating hero slides with image/video media, overlays, links, arrows, pagination and autoplay.",
"festival-calendar":"Displays dated festival cards and can mark upcoming or sale-enabled festivals.",
"section":"A flexible content area assembled from nested image, text, button, group, floating and custom Liquid blocks.",
"premium-collection-list":"Shows selected collections as a styled collection-navigation grid/list.",
"product-list":"Shows products from a selected collection, with a configurable heading area and product-card design.",
"slider-collection":"Shows manually selected collection cards in a responsive slider/grid presentation.",
"trust-badges":"Displays short reassurance messages with icons, such as secure payment or easy returns.",
"regional":"Displays editorial/blog content with a heading, supporting copy and configurable article cards.",
"testimonials-slider":"Displays customer testimonials in a navigable slider.",
"product-information":"The main product area: media gallery plus product details, variants, quantity and purchase actions.",
"product-recommendations":"Loads Shopify-generated related products after the page opens and displays them as product cards.",
"_blocks":"A general block container; this instance hosts the Judge.me review widget app block.",
"collection-banner":"Displays the collection heading/banner. Collection metafield images take priority over the Theme Editor fallback image.",
"collection-description":"Displays the collection description with an optional Read more / Read less expander.",
"promotional-collections":"Displays linked promotional collections supplied by a collection metafield.",
"main-collection":"Displays the collection product grid, filters, sorting, grid-density controls and pagination/infinite loading.",
}

SPECIAL={
"desktop_banner_image":"Fallback banner image used only when the collection's desktop banner metafield is empty.",
"mobile_image":"Mobile-specific media. In the Horizon slide it replaces desktop media on small screens; if absent, desktop media remains available.",
"desktop_image":"Primary image for desktop and the fallback when no mobile-specific image is supplied.",
"enable_overlay":"Adds the configured color overlay; opacity controls its strength.",
"overlay_opacity":"Controls overlay strength. Zero makes the overlay effectively invisible.",
"autoplay":"Automatically advances slides; the delay setting matters only when autoplay is enabled.",
"autoplay_delay":"Time between automatic slide changes; has no effect when autoplay is off.",
"show_arrows":"Shows or hides previous/next navigation controls.",
"show_pagination":"Shows or hides the slide position controls.",
"enable_infinite_scroll":"When enabled, more products load through the paginated-list behavior; when disabled, numbered pagination is rendered.",
"products_per_page":"Number of products requested per collection page.",
"enable_filtering":"Shows filters only when Shopify has filters available for this collection/search context.",
"enable_sorting":"Allows shoppers to change the Shopify collection sort order.",
"enable_grid_density":"Shows controls that let shoppers change product-card density.",
"enable_sticky_filter_sidebar":"Keeps the vertical filter area visible while scrolling; the top-offset setting controls where it stops.",
"max_visible_height":"Collapsed description height. Read-more controls appear only when the content exceeds this space.",
"collection":"Selects the Shopify collection that supplies cards/products.",
"blog":"Selects the blog that supplies article cards.",
"open_by_default":"Starts this accordion expanded. Shoppers can still collapse it.",
"metafield_key":"Chooses the key read from the product's custom metafield namespace.",
"gift_card_form":"Shows recipient fields only when the current product is a Shopify gift card.",
"show_pickup_availability":"Shows local pickup information when Shopify has availability data for the selected variant.",
"content_type":"Chooses image or video rendering; only the settings for the chosen media type affect the slide.",
"url":"Makes the whole Horizon slide a link when populated; blank leaves it non-clickable.",
"new_tab":"Opens the slide link in a new tab; relevant only when a URL is supplied.",
"link":"Destination for the related button/card. A blank destination may leave the element non-clickable or hide it where the implementation requires a link.",
"label":"Visible button wording. Keep it brief; blank labels normally produce no useful call to action.",
"image_ratio":"Controls the product-image frame and cropping behavior. Non-adapt ratios crop into the chosen frame; Adapt follows source-media proportions.",
}

MF={
"collection-banner":[("custom.collection_banner_desktop","Desktop collection banner; overrides the section fallback desktop image."),("custom.collection_banner_mobile","Mobile collection banner; used for mobile when populated; the desktop-resolved banner remains the fallback.")],
"promotional-collections":[("custom.promotional_collection","List of collections displayed by this section. Empty means there are no promotional collection items to show.")],
"family-products":[("custom.family_matches","Related family products used to build the family-product selector; empty suppresses its product choices.")],
"size-chart":[("custom.size_chart","Size-chart content/resource displayed by the size-chart block; empty prevents meaningful chart content.")],
"product-specifications":[("custom.structure_details, silhouette, movement_rating_new, sheer_opaque_new, stitching_type_new, neckline_new, sleeve_type_new, separate_sleeve_attached (fallback: seperate_sleeve_attached), closure_type_new, back_design_new, lining_details, fabric_gsm, fabric_type_new, fall_pico_new, structured_relevant_for_jackets, fit_note, dimensions_summary, lehenga_length, waist_dimension, included_items, component_fabrics, key_details","Populates the Product Specifications accordion; blank individual fields are omitted by the block's conditional output.")],
"craft-profile":[("custom.crafts, region, artisan_cluster_attribution","Populates craft and origin details; empty values leave no corresponding detail to show.")],
"occasion-festival":[("custom.[metafield_key]","Reads a merchant-selected key from the custom namespace, allowing this accordion to target a particular occasion/festival metafield.")],
"style-inspiration":[("custom.styling_notes, custom.design_artisan_story","Populates styling and artisan-story content.")],
"design-craftsmanship":[("custom.print_placement, print_technique_new, blouse_collar_new, strap_type_new, pattern_placement_new, embellishments_new, embellishment_placement_new, flare_details, drape_styles_new, pallu_design, border_type_new, border_width, color_primary_new, color_secondary_new, zari_type_new, weave_technique_new, waist_closure_new","Populates the Design & Craftsmanship accordion; only available product data is meaningful.")],
"care-profile":[("custom.care_profile","Populates product care instructions.")],
"claims-certifications":[("custom.claims_and_certifications, custom.gi_tag_status","Populates certification/claim information and GI-tag status.")],
"complete-the-look":[("custom.complete_the_look","Product list used by the multi-product Complete the Look interface; empty means no companion products." )],
"variant-picker":[("custom.name_for_bundle","Optional variant-option naming used by the variant picker for bundle products.")],
}

def setting_help(s):
    sid=s.get("id",""); typ=s.get("type","")
    if sid in SPECIAL: return SPECIAL[sid]
    human=title(s.get("label") or sid)
    if typ=="image_picker": return f"Selects the image used for {human.lower()}. Leaving it empty follows the section's coded fallback or shows no image. Exact dimensions are not enforced in the schema; use a consistent aspect ratio and sufficient resolution."
    if typ=="video": return "Selects Shopify-hosted video media. It is used only when the surrounding media/content choice calls for video."
    if typ in ("text","inline_richtext","richtext","textarea","html"): return f"Controls the visible {human.lower()}. Rich-text fields allow formatting; plain text fields do not. Blank text is omitted where the Liquid checks for blank content."
    if typ=="url": return f"Sets the destination for {human.lower()}. Leave blank when no navigation should occur."
    if typ=="checkbox": return f"Turns {human.lower()} on or off. Related settings may have no visible effect while this is off."
    if typ in ("range","number"): return f"Adjusts {human.lower()}. Use the displayed Theme Editor range; extreme values can noticeably change spacing, size, opacity or density."
    if typ in ("color","color_background"): return f"Sets {human.lower()}. A blank/dynamic palette value allows the relevant inherited/global color to apply."
    if typ=="select": return f"Chooses the {human.lower()} mode. Each option changes the corresponding layout or presentation described by its label."
    if typ in ("collection","product","blog","page"): return f"Selects the Shopify {typ} used as this setting's content source. Blank means there is no selected source."
    if typ=="font_picker": return f"Selects the font for {human.lower()}; this can override inherited typography for this element."
    return f"Controls {human.lower()} for this element. Change it in preview and check both desktop and mobile before saving."

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)

def add_toc(doc):
    p=doc.add_paragraph(); run=p.add_run(); fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'TOC \\o "1-3" \\h \\z \\u'); run._r.addnext(fld)

def note(doc,text,kind="NOTE"):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER; c=t.cell(0,0); shade(c,"FFF4D6" if kind=="IMPORTANT" else "EAF3F8"); c.paragraphs[0].add_run(kind+": ").bold=True; c.paragraphs[0].add_run(text)

def add_settings(doc, schema, current):
    settings=[x for x in schema.get("settings",[]) if x.get("type") not in ("header","paragraph") and x.get("id")]
    if not settings:
        doc.add_paragraph("This item has no merchant-facing settings in its schema.")
        return
    for s in settings:
        name=title(s.get("label") or s["id"])
        p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(7); p.paragraph_format.space_after=Pt(2)
        r=p.add_run(name); r.bold=True; r.font.color.rgb=RGBColor.from_string('6B1F2A')
        doc.add_paragraph(setting_help(s))
        p=doc.add_paragraph(style="List Bullet"); p.add_run("What to enter/select: ").bold=True
        if s.get("type")=="checkbox": p.add_run("Enable it only when you want the described feature visible or active.")
        elif s.get("options"): p.add_run("Choose the option that matches the required layout. Available choices: "+", ".join(title(o.get("label") or o.get("value")) for o in s["options"])+".")
        else: p.add_run("Use the Theme Editor control appropriate to this field and verify the live preview.")
        p=doc.add_paragraph(style="List Bullet"); p.add_run("Current template value: ").bold=True; p.add_run(fmt(current.get(s["id"], "Not explicitly saved (schema default applies)")))
        if "default" in s:
            p=doc.add_paragraph(style="List Bullet"); p.add_run("Schema default: ").bold=True; p.add_run(fmt(s["default"]))
        if "info" in s:
            p=doc.add_paragraph(style="List Bullet"); p.add_run("Theme note: ").bold=True; p.add_run(title(s["info"]) if str(s["info"]).startswith("t:") else str(s["info"]))

def main():
    doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.75); sec.right_margin=Inches(.75)
    styles=doc.styles
    styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(9.5)
    for sn,size,color in [('Title',28,'6B1F2A'),('Subtitle',16,'8C5A3C'),('Heading 1',20,'6B1F2A'),('Heading 2',15,'8C5A3C'),('Heading 3',11,'333333')]:
        st=styles[sn]; st.font.name='Aptos Display'; st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color)
    p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Shopify Theme Editor Management Guide')
    p=doc.add_paragraph(style='Subtitle'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Homepage, Product Page & Collection Page')
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Store administrator edition\n').bold=True; p.add_run(f'Prepared from the local theme code • {date.today().strftime("%d %B %Y")}')
    doc.add_page_break(); doc.add_heading('Table of Contents',0); add_toc(doc)
    doc.add_paragraph('If the table does not populate automatically, right-click it in Microsoft Word and choose Update Field → Update entire table.')
    doc.add_page_break(); doc.add_heading('1. Introduction',1)
    doc.add_paragraph('This guide is for the Shopify store administrator who manages page content and presentation through the visual Theme Editor. It covers the exact sections currently assigned to the Home page, Default product template and Default collection template in this theme.')
    doc.add_paragraph('The Theme Editor can manage section order, section and block settings, images, text, colors, spacing, selected products/collections/blogs, and enabled/disabled content. Product and collection metafield content is normally edited on the relevant product or collection record, not inside the Theme Editor.')
    note(doc,'Do not edit Liquid, JSON, JavaScript, CSS, app identifiers, custom Liquid, or code-like values unless a developer has reviewed the change. Content changes are stored as theme configuration; code changes alter how the theme works.', 'IMPORTANT')
    doc.add_heading('2. General Theme Editor Instructions',1)
    for x in ['In Shopify Admin, go to Online Store → Themes.','On the theme you intend to edit, choose Customize. Confirm that you are editing the correct live or draft theme.','Use the template selector at the top to choose Home page, Products → Default product, or Collections → Default collection.','Choose a representative product or collection when previewing a resource template. Different records can contain different metafields and images.','Select a section in the left sidebar. Change one setting at a time and watch the preview.','To reorder an allowed section or block, drag its handle. To remove one, open it and choose Remove. Use Add section or Add block only where the editor offers it. Static blocks cannot necessarily be removed or reordered.','Use the desktop/mobile preview icons to check both layouts.','Choose Save. Then preview the storefront in a new tab before publishing or leaving a draft theme.']:
        doc.add_paragraph(x,style='List Number')
    note(doc,'The three JSON templates are auto-generated by Shopify. Make structural changes through Customize, not by directly editing those JSON files.')

    quick=[]; metafields=[]
    for page,rel in TEMPLATES:
        d=load_template(rel); doc.add_page_break(); num={'Homepage':'3','Product Page':'4','Collection Page':'5'}[page]
        doc.add_heading(f'{num}. {page} Management',1)
        doc.add_paragraph(f"Template analyzed: {rel}. Sections are documented below in the exact saved order. Theme Editor drag handles determine whether a particular item can be reordered; static blocks are fixed by the theme.")
        for idx,sid in enumerate(d['order'],1):
            inst=d['sections'][sid]; typ=inst['type']; sf=section_file(typ); sch=schema_for(sf) if sf else {}
            display=title(sch.get('name') or inst.get('name') or typ)
            doc.add_heading(f'{display}',2)
            doc.add_heading('What this section does',3); doc.add_paragraph(PURPOSES.get(typ,f"Provides the theme's {display.lower()} presentation and behavior."))
            doc.add_heading('Where you will find it',3); doc.add_paragraph(f'Online Store → Themes → Customize → {page} → {display}')
            doc.add_paragraph(f"Saved position: {idx} of {len(d['order'])}. Section type: {typ}. " + ('This section is currently disabled.' if inst.get('disabled') else 'This section is currently enabled.'))
            doc.add_heading('Section settings',3); add_settings(doc,sch,inst.get('settings',{}))
            if typ in MF:
                doc.add_heading('Important metafield dependency',3)
                for k,v in MF[typ]: doc.add_paragraph(f'{k}: {v}',style='List Bullet'); metafields.append((page,display,k,v))
            quick.append((page,display,PURPOSES.get(typ,'Theme content section'),', '.join(title(k) for k in list(inst.get('settings',{}))[:5])))
            def walk(blocks, depth=0):
                for bid in (inst.get('block_order',list(blocks.keys())) if depth==0 else list(blocks.keys())):
                    if bid not in blocks: continue
                    b=blocks[bid]; bt=b['type']; bf=block_file(bt); bs=schema_for(bf) if bf else {}
                    bn=title(bs.get('name') or b.get('name') or ('Judge.me Review Widget' if bt.startswith('shopify://apps/judge-me') else bt))
                    doc.add_heading(f"Block: {bn}",3)
                    if bt.startswith('shopify://apps/'):
                        doc.add_paragraph('This is an app block supplied by Judge.me. Its availability and output depend on that app being installed, configured and able to load review data. Manage review content in Judge.me; use the Theme Editor only for this block’s exposed display settings.')
                    else:
                        doc.add_paragraph(f"This {('nested ' if depth else '')}block contributes {bn.lower()} content or layout inside {display}. It is currently {'disabled' if b.get('disabled') else 'enabled'}. Add/remove/reorder it only where the Theme Editor offers those controls; nested/static theme blocks may have restrictions.")
                    add_settings(doc,bs,b.get('settings',{}))
                    if bt in MF:
                        doc.add_heading('Block metafield dependency',3)
                        for k,v in MF[bt]: doc.add_paragraph(f'{k}: {v}',style='List Bullet'); metafields.append((page,f'{display} / {bn}',k,v))
                    if b.get('blocks'): walk(b['blocks'],depth+1)
            if inst.get('blocks'): doc.add_heading('Blocks',3); walk(inst['blocks'])

    doc.add_page_break(); doc.add_heading('6. Important Dependencies & Fallbacks',1)
    for s in [
        'Collection banner priority: custom.collection_banner_desktop is checked first. The section Desktop banner image is the fallback. custom.collection_banner_mobile supplies the mobile-specific banner. Without any resolved banner image, the section uses its minimal, no-image layout and switches block colors to the no-image color settings.',
        'Promotional Collections is data-driven: collection.metafields.custom.promotional_collection supplies the list. Theme Editor styling cannot create cards when that metafield is empty.',
        'Product detail accordions are data-driven: their headings, colors, icons and open/closed state are Theme Editor settings, but the actual facts come from product metafields. Blank metafields lead to omitted or empty detail content.',
        'Variant selection affects price, availability, media and purchase controls. The theme JavaScript updates the selected variant rather than treating these controls as static content.',
        'Product recommendations load asynchronously from Shopify. An empty result or request error means the recommendation area can remain hidden/empty despite its Theme Editor styling.',
        'Collection filters come from Shopify’s available collection filters. Enabling the theme control does not create filter definitions; configure those in Shopify Search & Discovery / store filtering settings.',
        'Theme palette references such as “Global palette: color8” inherit from global theme settings. Changing the global palette can alter several documented sections at once.',
        'Disabled blocks remain saved but do not render. Re-enable them in the editor before troubleshooting their individual settings.'
    ]: doc.add_paragraph(s,style='List Bullet')
    doc.add_heading('Metafield reference',2)
    t=doc.add_table(rows=1,cols=4); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(['Template','Section/block','Namespace and key','Purpose / empty behavior']): t.cell(0,i).text=h; shade(t.cell(0,i),'6B1F2A'); t.cell(0,i).paragraphs[0].runs[0].font.color.rgb=RGBColor(255,255,255)
    seen=set()
    for row in metafields:
        if row in seen: continue
        seen.add(row); cells=t.add_row().cells
        for i,v in enumerate(row): cells[i].text=v; cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP

    doc.add_heading('Settings to Change Carefully',2)
    for x in ['Custom Liquid block: it can output arbitrary Liquid/HTML and should be changed only with developer assistance.','Responsive/layout controls such as custom width/height, absolute/floating position, z-index-like layering, negative offsets, and large padding values can hide or overlap content.','Product metafield key in Occasion/Festival must match a real key in the custom namespace; a typo produces no content.','Infinite loading, sticky filters, gallery behavior, variant picker and purchase buttons rely on JavaScript and matching markup. Do not rename code-like identifiers or remove required static blocks.','App block identifiers belong to the app integration. Remove the Judge.me widget only if reviews are intentionally being removed from this template.']:
        doc.add_paragraph(x,style='List Bullet')

    doc.add_page_break(); doc.add_heading('7. Common Mistakes',1)
    mistakes=[
('Image does not appear','Check that the correct template and section are being edited, that the image field is populated, and that the section/block is enabled. For the collection banner, remember the collection metafield can override the section fallback.'),
('Mobile banner seems unchanged','Check the mobile preview and the relevant mobile image. On collection pages, custom.collection_banner_mobile is the resource-specific source; on Horizon slides, the mobile image is used only when supplied.'),
('Button is missing or does nothing','Check both the label and destination, and confirm the containing block/group is enabled. A blank URL intentionally leaves a Horizon slide non-clickable.'),
('Product accordion has a heading but no facts','Add the required custom product metafields to the previewed product. Theme Editor styling does not supply the underlying product facts.'),
('Variant, size or swatch options are missing','Confirm the product has those Shopify options/variants and that they are available. The picker cannot manufacture options that do not exist in product data.'),
('Recommendations are empty','Shopify may have returned no recommendations, or the asynchronous request may have failed. Test another product and confirm recommendation data/settings.'),
('Collection filters are absent','Confirm filtering is enabled in this block and that Shopify has configured filters applicable to the collection.'),
('Promotional collection circles/cards are empty','Populate collection.metafields.custom.promotional_collection on the collection being previewed.'),
('Setting has no visible effect','Look for a parent toggle, content-type choice, blank content source, disabled parent block, responsive-only setting, or global palette inheritance.'),
('Changes appear on the wrong products/collections','Verify the template assignment on the resource and the resource selected in the Theme Editor preview.')]
    for a,b in mistakes: doc.add_heading(a,2); doc.add_paragraph(b)

    doc.add_page_break(); doc.add_heading('8. Quick Reference',1)
    t=doc.add_table(rows=1,cols=4); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(['Template','Section','Main purpose','Important current settings']): t.cell(0,i).text=h; shade(t.cell(0,i),'6B1F2A'); t.cell(0,i).paragraphs[0].runs[0].font.color.rgb=RGBColor(255,255,255)
    for row in quick:
        cells=t.add_row().cells
        for i,v in enumerate(row): cells[i].text=v
    doc.add_paragraph('End of guide. Always preview a representative product and collection on both desktop and mobile before publishing theme changes.')
    doc.core_properties.title='Shopify Theme Editor Management Guide'
    doc.core_properties.subject='Homepage, Product Page & Collection Page'
    doc.core_properties.author='Theme documentation generated from the local Shopify theme codebase'
    doc.save(OUT)
    print(OUT)

if __name__=='__main__': main()
